"""Verify deterministic ETL specification content and file renderers."""

import base64
import re
import unittest
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import yaml
from docx import Document

from agent.etl_specification import (
    _PROJECT_PURPOSE,
    _mapping_figure_images,
    build_etl_specification,
    build_etl_specification_bundle,
    build_etl_specification_document,
)
from agent.validation import validate_spec_contents, validate_specs


class EtlSpecificationTest(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.root = Path(__file__).resolve().parents[1]
        cls.specs = validate_specs("person", cls.root / "specs")

    def test_document_uses_target_order_and_explicit_nulls(self):
        document = build_etl_specification_document(self.specs)

        self.assertEqual(document.target_table, "person")
        self.assertEqual(document.cdm_version, "5.4")
        self.assertEqual(
            [row.destination_field for row in document.rows],
            [field.name for field in self.specs.target_schema.fields],
        )
        person_id = document.rows[0]
        self.assertEqual(
            person_id.source_field,
            "cai_01_patient.patient_id",
        )
        self.assertIn("Cast patient_id", person_id.logic)
        location = next(
            row
            for row in document.rows
            if row.destination_field == "location_id"
        )
        self.assertEqual(location.source_field, "—")
        self.assertIn("Set NULL", location.logic)

    def test_mapping_notes_comments_and_change_log_are_documented(self):
        mapping = yaml.safe_load(
            (self.root / "specs" / "mappings" / "person.yml").read_text(
                encoding="utf-8"
            )
        )
        mapping["notes"] = "One row per reconciled source patient."
        mapping["fields"][0]["comment"] = "Preserve source lineage."
        mapping["change_log"] = [
            {
                "date": "2026-08-04",
                "description": "Added ETL documentation metadata.",
                "author": "Data team",
            }
        ]
        specs = validate_spec_contents(
            "person",
            yaml.safe_dump(mapping, sort_keys=False),
            {
                "cai_01_patient.yml": (
                    self.root
                    / "specs"
                    / "source_schema"
                    / "cai_01_patient.yml"
                ).read_text(encoding="utf-8")
            },
            (
                self.root / "specs" / "target_schema" / "person.yml"
            ).read_text(encoding="utf-8"),
        )

        document = build_etl_specification_document(specs)

        self.assertEqual(
            document.table_notes,
            "One row per reconciled source patient.",
        )
        self.assertIn("Preserve source lineage.", document.rows[0].comment)
        self.assertEqual(
            document.changes,
            ((
                "2026-08-04",
                "Added ETL documentation metadata.",
                "Data team",
            ),),
        )

    def test_renders_markdown_word_and_pdf_without_ai(self):
        project = {
            "project_name": "Harmonia OMOP",
            "project_description": "Maps project source data into OMOP CDM.",
        }
        markdown = build_etl_specification(self.specs, "md", **project)
        word = build_etl_specification(self.specs, "docx", **project)
        pdf = build_etl_specification(self.specs, "pdf", **project)

        self.assertEqual(markdown.file_name, "person_etl_specification.md")
        markdown_text = markdown.content.decode("utf-8")
        self.assertTrue(markdown_text.startswith("# Harmonia OMOP\n"))
        self.assertIn(_PROJECT_PURPOSE, markdown_text)
        self.assertNotIn("```mermaid", markdown_text)
        image_match = re.search(
            r"data:image/png;base64,([^\)]+)",
            markdown_text,
        )
        self.assertIsNotNone(image_match)
        self.assertEqual(
            base64.b64decode(image_match.group(1)),
            _mapping_figure_images(self.specs)[0],
        )
        self.assertIn(
            "| Destination Field | Source field | Logic | Comment field |",
            markdown_text,
        )
        self.assertIn("### Change log", markdown_text)

        self.assertTrue(word.content.startswith(b"PK\x03\x04"))
        word_document = Document(BytesIO(word.content))
        self.assertEqual(
            word_document.paragraphs[0].text,
            "Harmonia OMOP",
        )
        self.assertIn(
            "person ETL specification",
            [paragraph.text for paragraph in word_document.paragraphs],
        )
        self.assertEqual(len(word_document.sections), 2)
        self.assertLess(
            word_document.sections[0].page_width,
            word_document.sections[0].page_height,
        )
        self.assertGreater(
            word_document.sections[1].page_width,
            word_document.sections[1].page_height,
        )
        self.assertEqual(len(word_document.inline_shapes), 1)
        self.assertGreaterEqual(len(word_document.tables), 2)
        self.assertEqual(
            [cell.text for cell in word_document.tables[0].rows[0].cells],
            [
                "Destination Field",
                "Source field",
                "Logic",
                "Comment field",
            ],
        )
        self.assertEqual(
            word.content,
            build_etl_specification(
                self.specs,
                "docx",
                **project,
            ).content,
        )
        with ZipFile(BytesIO(word.content)) as archive:
            document_xml = archive.read("word/document.xml")
            footer_xml = archive.read("word/footer1.xml")
            word_figure = archive.read("word/media/image1.png")
        self.assertEqual(word_figure, _mapping_figure_images(self.specs)[0])
        self.assertIn(b"w:tblHeader", document_xml)
        self.assertIn(b"w:cantSplit", document_xml)
        self.assertIn(b'<w:top w:w="160" w:type="dxa"', document_xml)
        self.assertIn(b"PAGE", footer_xml)
        self.assertIn(b"NUMPAGES", footer_xml)

        self.assertTrue(pdf.content.startswith(b"%PDF"))
        media_boxes = re.findall(
            rb"/MediaBox\s*\[\s*0\s+0\s+([0-9.]+)\s+([0-9.]+)\s*\]",
            pdf.content,
        )
        self.assertTrue(
            any(float(width) < float(height) for width, height in media_boxes)
        )
        self.assertTrue(
            any(float(width) > float(height) for width, height in media_boxes)
        )
        self.assertEqual(
            pdf.content,
            build_etl_specification(
                self.specs,
                "pdf",
                **project,
            ).content,
        )

    def test_combines_tables_in_one_document_deterministically(self):
        visit_specs = validate_specs(
            "visit_occurrence",
            self.root / "specs",
        )

        project = {
            "project_name": "Harmonia OMOP",
            "project_description": "Maps project source data into OMOP CDM.",
        }
        markdown = build_etl_specification_bundle(
            [self.specs, visit_specs], "md", **project
        )
        word = build_etl_specification_bundle(
            [self.specs, visit_specs], "docx", **project
        )
        pdf = build_etl_specification_bundle(
            [self.specs, visit_specs], "pdf", **project
        )

        self.assertEqual(
            markdown.media_type,
            "text/markdown; charset=utf-8",
        )
        self.assertEqual(markdown.file_name, "omop_etl_specification.md")
        markdown_text = markdown.content.decode("utf-8")
        self.assertTrue(markdown_text.startswith("# Harmonia OMOP\n"))
        self.assertEqual(markdown_text.count("## OMOP ETL specification"), 1)
        self.assertIn("## person ETL specification", markdown_text)
        self.assertIn("## visit_occurrence ETL specification", markdown_text)

        self.assertEqual(
            word.media_type,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        self.assertEqual(word.file_name, "omop_etl_specification.docx")
        word_document = Document(BytesIO(word.content))
        self.assertEqual(len(word_document.inline_shapes), 2)
        self.assertEqual(len(word_document.sections), 4)
        self.assertIn(
            "person ETL specification",
            [paragraph.text for paragraph in word_document.paragraphs],
        )
        self.assertIn(
            "visit_occurrence ETL specification",
            [paragraph.text for paragraph in word_document.paragraphs],
        )

        self.assertEqual(pdf.media_type, "application/pdf")
        self.assertEqual(pdf.file_name, "omop_etl_specification.pdf")
        self.assertTrue(pdf.content.startswith(b"%PDF"))

        for output_format, artifact in (
            ("md", markdown),
            ("docx", word),
            ("pdf", pdf),
        ):
            self.assertEqual(
                artifact.content,
                build_etl_specification_bundle(
                    [visit_specs, self.specs],
                    output_format,
                    **project,
                ).content,
            )


if __name__ == "__main__":
    unittest.main()
