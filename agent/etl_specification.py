"""Build deterministic ETL specification documents from validated mappings."""

from __future__ import annotations

from base64 import b64encode
from dataclasses import dataclass
from html import escape as html_escape
from io import BytesIO
from typing import Literal
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    BaseDocTemplate,
    CondPageBreak,
    Frame,
    Image as PdfImage,
    KeepTogether,
    LongTable,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    TableStyle,
)

from .contracts import FieldMapping, TargetField
from .validation import ValidatedSpecs

EtlSpecificationFormat = Literal["md", "docx", "pdf"]

_GREEN = "0B6657"
_GREEN_SOFT = "DCEEE8"
_INK = "081917"
_MUTED = "536B65"
_LINE = "DCE5E1"
_AMBER = "C56F45"
_MAXIMUM_ROWS = 500
_DEFAULT_PROJECT_NAME = "OMOP ETL project"
_PROJECT_PURPOSE = (
    "This document records the requirements, assumptions, source-to-target "
    "mappings, business rules and transformations used to convert project "
    "source data into the OMOP Common Data Model. It is intended as an "
    "implementation and review reference for data engineers, clinical and "
    "terminology experts, and researchers."
)


@dataclass(frozen=True)
class EtlSpecificationRow:
    """One field-level row in an ETL specification."""

    destination_field: str
    source_field: str
    logic: str
    comment: str


@dataclass(frozen=True)
class EtlSpecificationDocument:
    """Renderer-independent ETL specification content."""

    cdm_version: str
    target_table: str
    table_description: str
    table_notes: str
    source_models: tuple[str, ...]
    relationships: tuple[str, ...]
    rows: tuple[EtlSpecificationRow, ...]
    changes: tuple[tuple[str, str, str], ...]


@dataclass(frozen=True)
class EtlSpecificationArtifact:
    """One bounded downloadable ETL specification file."""

    content: bytes
    file_name: str
    media_type: str


def _source_label(mapping: FieldMapping) -> str:
    if not mapping.source_fields:
        return "—"
    return "\n".join(
        f"{reference.model}.{reference.field}"
        for reference in mapping.source_fields
    )


def _logic(mapping: FieldMapping, target: TargetField) -> str:
    transformation = mapping.transformation.strip()
    if mapping.action == "null":
        return transformation or "Set NULL."

    parts: list[str] = []
    if mapping.mapping_table_name:
        parts.append(f"Lookup using {mapping.mapping_table_name}.")
    if transformation:
        parts.append(transformation)
    elif mapping.source_fields:
        parts.append(
            "Direct mapping; conform the value to the OMOP "
            f"{target.data_type} datatype."
        )
    return " ".join(parts) or "No transformation logic recorded."


def _comment(mapping: FieldMapping) -> str:
    parts = []
    if mapping.comment.strip():
        parts.append(mapping.comment.strip())
    if mapping.review_comment and mapping.review_comment.strip():
        parts.append(f"Review: {mapping.review_comment.strip()}")
    return "\n".join(parts) or "—"


def _relationships(specs: ValidatedSpecs) -> tuple[str, ...]:
    relationships = [
        (
            f"{join.join_type.upper()} JOIN "
            f"{join.left.model}.{join.left.field} = "
            f"{join.right.model}.{join.right.field}"
        )
        for join in specs.mapping.joins
    ]
    if specs.mapping.union_all:
        relationships.append(
            "UNION ALL: " + ", ".join(specs.mapping.union_all)
        )
    return tuple(relationships)


def build_etl_specification_document(
    specs: ValidatedSpecs,
) -> EtlSpecificationDocument:
    """Create a deterministic document model in OMOP target-field order."""
    mappings = {
        mapping.target_field: mapping
        for mapping in specs.mapping.fields
    }
    rows: list[EtlSpecificationRow] = []
    for target in specs.target_schema.fields:
        mapping = mappings.get(target.name)
        if mapping is None:
            rows.append(
                EtlSpecificationRow(
                    destination_field=target.name,
                    source_field="—",
                    logic=(
                        "Set NULL because this optional OMOP field is not "
                        "configured in the mapping."
                    ),
                    comment="—",
                )
            )
            continue
        rows.append(
            EtlSpecificationRow(
                destination_field=target.name,
                source_field=_source_label(mapping),
                logic=_logic(mapping, target),
                comment=_comment(mapping),
            )
        )

    if len(rows) > _MAXIMUM_ROWS:
        raise ValueError("The ETL specification exceeds the row limit.")

    return EtlSpecificationDocument(
        cdm_version=specs.target_schema.cdm_version,
        target_table=specs.mapping.target_table,
        table_description=specs.target_schema.description.strip(),
        table_notes=specs.mapping.notes.strip(),
        source_models=tuple(specs.mapping.source_models),
        relationships=_relationships(specs),
        rows=tuple(rows),
        changes=tuple(
            (
                change.date.isoformat(),
                change.description.strip(),
                change.author.strip() or "—",
            )
            for change in specs.mapping.change_log
        ),
    )


def _markdown_cell(value: str) -> str:
    return (
        html_escape(value, quote=False)
        .replace("|", "\\|")
        .replace("\n", "<br>")
    )


def _markdown_text(value: str) -> str:
    """Prevent submitted specification text from becoming active HTML."""
    return html_escape(value, quote=False)


def render_markdown(
    document: EtlSpecificationDocument,
    specs: ValidatedSpecs,
) -> bytes:
    """Render a self-contained Markdown ETL specification."""
    figure = _mapping_figure_images(specs)[0]
    encoded_figure = b64encode(figure).decode("ascii")
    lines = [
        f"# {document.target_table} ETL specification",
        "",
        f"**OMOP CDM version:** {document.cdm_version}",
        f"**Source models:** {', '.join(document.source_models)}",
        "**Generation method:** Deterministic; no AI request",
        "",
        "## Overview",
        "",
        _markdown_text(
            document.table_notes
            or document.table_description
            or "No table-level notes recorded."
        ),
        "",
        "## Mapping overview",
        "",
        "![Source-to-OMOP mapping grid]"
        f"(data:image/png;base64,{encoded_figure})",
        "",
    ]
    if document.relationships:
        lines.extend(["## Source relationships", ""])
        lines.extend(
            f"- {_markdown_text(item)}"
            for item in document.relationships
        )
        lines.append("")
    lines.extend(
        [
            "## Field mapping",
            "",
            "| Destination Field | Source field | Logic | Comment field |",
            "|---|---|---|---|",
        ]
    )
    lines.extend(
        "| "
        + " | ".join(
            _markdown_cell(value)
            for value in (
                row.destination_field,
                row.source_field,
                row.logic,
                row.comment,
            )
        )
        + " |"
        for row in document.rows
    )
    lines.extend(
        [
            "",
            "## Change log",
            "",
            "| Date | Description | Author |",
            "|---|---|---|",
        ]
    )
    if document.changes:
        lines.extend(
            "| "
            + " | ".join(_markdown_cell(value) for value in change)
            + " |"
            for change in document.changes
        )
    else:
        lines.append("| — | No changes recorded. | — |")
    lines.append("")
    return "\n".join(lines).encode("utf-8")


def render_markdown_collection(
    entries: tuple[tuple[EtlSpecificationDocument, ValidatedSpecs], ...],
    *,
    project_name: str,
    project_description: str,
) -> bytes:
    """Render multiple OMOP table sections in one Markdown document."""
    versions = sorted({document.cdm_version for document, _ in entries})
    tables = [document.target_table for document, _ in entries]
    lines = [
        f"# {_markdown_text(project_name)}",
        "",
        "## OMOP ETL specification",
        "",
        _markdown_text(
            project_description
            or "No project description was provided."
        ),
        "",
        "### Purpose",
        "",
        _PROJECT_PURPOSE,
        "",
        "### Document scope",
        "",
        f"**OMOP CDM version:** {', '.join(versions)}",
        f"**Included OMOP tables:** {len(tables)}",
        "",
        "**Generation method:** Deterministic; no AI request",
        "",
        "### Included tables",
        "",
        *(
            f"- {_markdown_text(document.target_table)}"
            for document, _ in entries
        ),
        "",
        "---",
        "",
    ]
    for document, specs in entries:
        table_lines = render_markdown(document, specs).decode("utf-8").splitlines()
        for line in table_lines:
            if line.startswith("## "):
                lines.append(f"### {line[3:]}")
            elif line.startswith("# "):
                lines.append(f"## {line[2:]}")
            else:
                lines.append(line)
    return "\n".join(lines).encode("utf-8")


def _font(size: int, bold: bool = False):
    name = "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf"
    try:
        return ImageFont.truetype(name, size=size)
    except OSError:
        return ImageFont.load_default(size=size)


def _mapping_status(mapping: FieldMapping | None) -> str:
    if mapping is None:
        return "unmapped"
    if mapping.action == "null":
        return "null"
    if mapping.mapping_table_name:
        return "lookup"
    return "mapped"


def _draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str,
) -> None:
    """Draw a mapping-grid style orthogonal connector and arrowhead."""
    middle = start[0] + max(40, (end[0] - start[0]) // 2)
    points = [start, (middle, start[1]), (middle, end[1]), end]
    draw.line(points, fill=color, width=4, joint="curve")
    draw.polygon(
        [end, (end[0] - 14, end[1] - 8), (end[0] - 14, end[1] + 8)],
        fill=color,
    )


def _draw_dashed_line(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[int, int]],
    color: str,
) -> None:
    for start, end in zip(points, points[1:]):
        length = abs(end[0] - start[0]) + abs(end[1] - start[1])
        if not length:
            continue
        horizontal = start[1] == end[1]
        for offset in range(0, length, 18):
            finish = min(offset + 10, length)
            if horizontal:
                segment = (
                    (start[0] + offset, start[1]),
                    (start[0] + finish, start[1]),
                )
            else:
                direction = 1 if end[1] >= start[1] else -1
                segment = (
                    (start[0], start[1] + direction * offset),
                    (start[0], start[1] + direction * finish),
                )
            draw.line(segment, fill=color, width=3)


def _draw_mapping_card(
    draw: ImageDraw.ImageDraw,
    *,
    box: tuple[int, int, int, int],
    kind: str,
    title: str,
    subtitle: str,
    rows: list[tuple[str, str, str]],
    header_fill: str,
    border: str,
    handles: Literal["both", "input", "output"],
) -> dict[str, tuple[int, int]]:
    """Draw one static equivalent of a mapping-grid React Flow tile."""
    left, top, right, bottom = box
    row_height = 62
    header_height = 86
    draw.rounded_rectangle(
        box,
        radius=14,
        fill="#FFFFFF",
        outline=border,
        width=3,
    )
    draw.rounded_rectangle(
        (left, top, right, top + header_height),
        radius=14,
        fill=header_fill,
    )
    draw.rectangle(
        (left, top + header_height - 14, right, top + header_height),
        fill=header_fill,
    )
    draw.line(
        (left, top + header_height, right, top + header_height),
        fill=f"#{_LINE}",
        width=2,
    )
    draw.text(
        (left + 18, top + 12),
        kind.upper(),
        fill=f"#{_MUTED}",
        font=_font(14, bold=True),
    )
    title_size = 22
    title_font = _font(title_size, bold=True)
    while (
        title_size > 14
        and draw.textlength(title, font=title_font) > right - left - 36
    ):
        title_size -= 1
        title_font = _font(title_size, bold=True)
    draw.text(
        (left + 18, top + 33),
        title,
        fill=f"#{_INK}",
        font=title_font,
    )
    draw.text(
        (left + 18, top + 62),
        subtitle,
        fill=f"#{_MUTED}",
        font=_font(14),
    )

    anchors: dict[str, tuple[int, int]] = {}
    stripe_colors = {
        "lookup": f"#{_AMBER}",
        "mapped": "#356D73",
        "null": "#AEBBB6",
        "unmapped": f"#{_LINE}",
    }
    for index, (row_id, label, meta) in enumerate(rows):
        row_top = top + header_height + index * row_height
        row_bottom = row_top + row_height
        draw.rectangle(
            (left + 2, row_top, right - 2, row_bottom),
            fill="#FFFFFF" if index % 2 == 0 else "#F8FAF9",
        )
        status = row_id.split("|", 1)[0]
        actual_id = row_id.split("|", 1)[-1]
        draw.rectangle(
            (left + 2, row_top, left + 8, row_bottom),
            fill=stripe_colors.get(status, "#356D73"),
        )
        draw.text(
            (left + 20, row_top + 10),
            label,
            fill=f"#{_INK}",
            font=_font(18, bold=True),
        )
        draw.text(
            (left + 20, row_top + 34),
            meta,
            fill=f"#{_MUTED}",
            font=_font(14),
        )
        if index < len(rows) - 1:
            draw.line(
                (left, row_bottom, right, row_bottom),
                fill=f"#{_LINE}",
                width=2,
            )
        anchor_y = row_top + row_height // 2
        anchors[actual_id] = (right, anchor_y)
        if handles in {"both", "input"}:
            draw.ellipse(
                (left - 6, anchor_y - 6, left + 6, anchor_y + 6),
                fill="#FFFFFF",
                outline=border,
                width=3,
            )
        if handles in {"both", "output"}:
            draw.ellipse(
                (right - 6, anchor_y - 6, right + 6, anchor_y + 6),
                fill="#FFFFFF",
                outline=border,
                width=3,
            )
    return anchors


def _mapping_figure_images(specs: ValidatedSpecs) -> list[bytes]:
    """Render one complete mapping-grid image for every output format."""
    chunks = [list(specs.target_schema.fields)]
    mappings = {
        mapping.target_field: mapping
        for mapping in specs.mapping.fields
    }
    images: list[bytes] = []
    for chunk_index, chunk in enumerate(chunks):
        width = 1800
        card_header_height = 86
        row_height = 62
        top = 160
        source_width = 500
        lookup_width = 390
        target_width = 500
        source_x = 42
        lookup_x = 705
        target_x = 1258

        chunk_names = {target.name for target in chunk}
        relevant_fields: dict[str, set[str]] = {
            model: set() for model in specs.mapping.source_models
        }
        for target_name in chunk_names:
            mapping = mappings.get(target_name)
            if mapping:
                for reference in mapping.source_fields:
                    relevant_fields[reference.model].add(reference.field)
        if chunk_index == 0:
            for join in specs.mapping.joins:
                relevant_fields[join.left.model].add(join.left.field)
                relevant_fields[join.right.model].add(join.right.field)

        source_rows: dict[str, list[tuple[str, str, str]]] = {}
        source_height = 0
        for model_name in specs.mapping.source_models:
            model = specs.source_models[model_name]
            rows = [
                (
                    f"mapped|{column.name}",
                    column.name,
                    column.data_type or "type not declared",
                )
                for column in model.columns
                if column.name in relevant_fields[model_name]
            ]
            source_rows[model_name] = rows
            source_height += card_header_height + max(1, len(rows)) * row_height
        source_height += max(0, len(source_rows) - 1) * 30

        lookup_fields = [
            target
            for target in chunk
            if mappings.get(target.name)
            and mappings[target.name].mapping_table_name
        ]
        lookup_height = len(lookup_fields) * (card_header_height + row_height + 22)
        target_height = card_header_height + max(1, len(chunk)) * row_height
        content_height = max(source_height, lookup_height, target_height)
        height = top + content_height + 62
        image = Image.new("RGB", (width, height), "#FFFFFF")
        draw = ImageDraw.Draw(image)
        for dot_x in range(18, width, 28):
            for dot_y in range(132, height, 28):
                draw.ellipse(
                    (dot_x, dot_y, dot_x + 2, dot_y + 2),
                    fill="#D5DFDB",
                )
        title_suffix = (
            f" · {chunk_index + 1}/{len(chunks)}"
            if len(chunks) > 1
            else ""
        )
        draw.text(
            (42, 24),
            f"{specs.mapping.target_table} source-to-OMOP flow{title_suffix}",
            fill=f"#{_INK}",
            font=_font(31, bold=True),
        )
        legend = [
            (f"#{_GREEN}", "Source table"),
            ("#356D73", "Derive / transform"),
            (f"#{_AMBER}", "Lookup table"),
            ("#AEBBB6", "Set NULL"),
            (f"#{_LINE}", "Not configured"),
        ]
        legend_x = 42
        for color, label in legend:
            draw.ellipse(
                (legend_x, 88, legend_x + 16, 104),
                fill=color,
            )
            draw.text(
                (legend_x + 24, 84),
                label,
                fill=f"#{_MUTED}",
                font=_font(15),
            )
            legend_x += 170

        source_boxes: dict[str, tuple[int, int, int, int]] = {}
        source_anchor_positions: dict[str, tuple[int, int]] = {}
        next_source_y = top
        for model_name in specs.mapping.source_models:
            rows = source_rows[model_name]
            model = specs.source_models[model_name]
            card_height = card_header_height + max(1, len(rows)) * row_height
            box = (
                source_x,
                next_source_y,
                source_x + source_width,
                next_source_y + card_height,
            )
            source_boxes[model_name] = box
            if rows:
                anchors = _draw_mapping_card(
                    draw,
                    box=box,
                    kind="source",
                    title=model_name,
                    subtitle=f"{len(model.columns)} available fields",
                    rows=rows,
                    header_fill=f"#{_GREEN_SOFT}",
                    border=f"#{_GREEN}",
                    handles="both",
                )
                source_anchor_positions.update(
                    {
                        f"{model_name}.{field}": position
                        for field, position in anchors.items()
                    }
                )
            else:
                _draw_mapping_card(
                    draw,
                    box=box,
                    kind="source",
                    title=model_name,
                    subtitle="No fields used on this page",
                    rows=[("unmapped|empty", "No mapped fields", "")],
                    header_fill=f"#{_GREEN_SOFT}",
                    border=f"#{_GREEN}",
                    handles="both",
                )
            next_source_y += card_height + 30

        target_rows: list[tuple[str, str, str]] = []
        for target in chunk:
            mapping = mappings.get(target.name)
            status_name = _mapping_status(mapping)
            status_label = {
                "lookup": "Lookup",
                "mapped": "Transform",
                "null": "NULL",
                "unmapped": "Not configured",
            }[status_name]
            requirement = "Required" if target.required else "Optional"
            target_rows.append(
                (
                    f"{status_name}|{target.name}",
                    target.name,
                    f"{target.data_type} · {requirement} · {status_label}",
                )
            )
        target_box = (
            target_x,
            top,
            target_x + target_width,
            top + target_height,
        )
        target_anchors = _draw_mapping_card(
            draw,
            box=target_box,
            kind="target",
            title=specs.mapping.target_table,
            subtitle=(
                f"{len(specs.mapping.fields)} of "
                f"{len(specs.target_schema.fields)} configured"
            ),
            rows=target_rows,
            header_fill="#E4EFF0",
            border="#B7D2D5",
            handles="input",
        )

        lookup_anchors: dict[str, tuple[tuple[int, int], tuple[int, int]]] = {}
        lookup_boxes: list[tuple[int, int, int, int]] = []
        next_lookup_y = top
        for target in lookup_fields:
            mapping = mappings[target.name]
            card_height = card_header_height + row_height
            box = (
                lookup_x,
                next_lookup_y,
                lookup_x + lookup_width,
                next_lookup_y + card_height,
            )
            lookup_boxes.append(box)
            anchors = _draw_mapping_card(
                draw,
                box=box,
                kind="lookup",
                title=mapping.mapping_table_name or "",
                subtitle="Lookup table",
                rows=[(
                    f"lookup|{target.name}",
                    target.name,
                    "Controlled lookup",
                )],
                header_fill="#FBE9DF",
                border="#EDC4AB",
                handles="both",
            )
            y = anchors[target.name][1]
            lookup_anchors[target.name] = (
                (lookup_x, y),
                (lookup_x + lookup_width, y),
            )
            next_lookup_y += card_height + 22

        # Draw connectors on a transparent layer and clip them beneath cards,
        # matching React Flow's edge-behind-node behaviour.
        edge_layer = Image.new("RGBA", (width, height), (0, 0, 0, 0))
        edge_draw = ImageDraw.Draw(edge_layer)
        for target in chunk:
            mapping = mappings.get(target.name)
            if not mapping or mapping.action == "null":
                continue
            destination = (target_x, target_anchors[target.name][1])
            if mapping.mapping_table_name:
                lookup_input, lookup_output = lookup_anchors[target.name]
                _draw_arrow(
                    edge_draw,
                    lookup_output,
                    destination,
                    f"#{_AMBER}",
                )
                destination = lookup_input
                edge_color = f"#{_AMBER}"
            else:
                edge_color = "#356D73"
            for reference in mapping.source_fields:
                source = source_anchor_positions.get(
                    f"{reference.model}.{reference.field}"
                )
                if source:
                    _draw_arrow(edge_draw, source, destination, edge_color)

        if chunk_index == 0:
            for join in specs.mapping.joins:
                left = source_anchor_positions.get(
                    f"{join.left.model}.{join.left.field}"
                )
                right = source_anchor_positions.get(
                    f"{join.right.model}.{join.right.field}"
                )
                if left and right:
                    outside_x = source_x - 20
                    _draw_dashed_line(
                        edge_draw,
                        [
                            (source_x, left[1]),
                            (outside_x, left[1]),
                            (outside_x, right[1]),
                            (source_x, right[1]),
                        ],
                        f"#{_GREEN}",
                    )

        for left, card_top, right, card_bottom in (
            [*source_boxes.values(), target_box, *lookup_boxes]
        ):
            edge_draw.rectangle(
                (left + 2, card_top + 2, right - 2, card_bottom - 2),
                fill=(0, 0, 0, 0),
            )
        image.paste(edge_layer, (0, 0), edge_layer)

        # Restore connector endpoints above the clipped edge layer.
        for position in source_anchor_positions.values():
            draw.ellipse(
                (position[0] - 6, position[1] - 6, position[0] + 6, position[1] + 6),
                fill="#FFFFFF",
                outline=f"#{_GREEN}",
                width=3,
            )
        for target_name, position in target_anchors.items():
            draw.ellipse(
                (target_x - 6, position[1] - 6, target_x + 6, position[1] + 6),
                fill="#FFFFFF",
                outline=(
                    f"#{_AMBER}"
                    if _mapping_status(mappings.get(target_name)) == "lookup"
                    else "#356D73"
                ),
                width=3,
            )
        buffer = BytesIO()
        image.save(buffer, format="PNG", optimize=True)
        images.append(buffer.getvalue())
    return images


def _shade_docx_cell(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_docx_cell_text(cell, value: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    for index, line in enumerate(value.splitlines() or [""]):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(_INK)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _set_docx_cell_margins(
    cell,
    *,
    top: int = 80,
    bottom: int = 80,
    start: int = 120,
    end: int = 120,
) -> None:
    """Apply explicit DXA cell padding without imposing a fixed row height."""
    properties = cell._tc.get_or_add_tcPr()
    existing = properties.find(qn("w:tcMar"))
    if existing is not None:
        properties.remove(existing)
    margins = OxmlElement("w:tcMar")
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        margin = OxmlElement(f"w:{edge}")
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    properties.append(margins)


def _add_docx_heading(document, text: str, level: int) -> None:
    heading = document.add_heading(text, level=level)
    heading.style.font.name = "Arial"
    heading.style.font.color.rgb = RGBColor.from_string(_GREEN)
    heading.paragraph_format.keep_with_next = True


def _keep_docx_row_together(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def _add_docx_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_name
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def _add_docx_footer(document, target_table: str) -> None:
    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"{target_table} ETL specification  ·  Page ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(_MUTED)
    _add_docx_page_field(footer, "PAGE")
    footer.add_run(" of ")
    _add_docx_page_field(footer, "NUMPAGES")


def _normalize_docx_archive(content: bytes) -> bytes:
    """Remove ZIP timestamps so identical inputs produce identical bytes."""
    source_buffer = BytesIO(content)
    output_buffer = BytesIO()
    with ZipFile(source_buffer, "r") as source, ZipFile(
        output_buffer,
        "w",
        compression=ZIP_DEFLATED,
        compresslevel=9,
    ) as destination:
        for name in sorted(source.namelist()):
            entry = ZipInfo(name, date_time=(1980, 1, 1, 0, 0, 0))
            entry.compress_type = ZIP_DEFLATED
            entry.external_attr = 0o644 << 16
            destination.writestr(entry, source.read(name))
    return output_buffer.getvalue()


def _configure_docx_section(section, orientation: WD_ORIENT) -> None:
    """Apply the stable A4 geometry for one Word section."""
    section.orientation = orientation
    if orientation == WD_ORIENT.PORTRAIT:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    else:
        section.page_width = Mm(297)
        section.page_height = Mm(210)
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)
    section.left_margin = Mm(12)
    section.right_margin = Mm(12)


def _append_docx_specification(
    document,
    content: EtlSpecificationDocument,
    specs: ValidatedSpecs,
    *,
    first: bool,
) -> None:
    """Append one portrait overview and landscape table section."""
    if first:
        portrait_section = document.sections[0]
    else:
        portrait_section = document.add_section(WD_SECTION.NEW_PAGE)
        portrait_section.footer.is_linked_to_previous = True
    _configure_docx_section(portrait_section, WD_ORIENT.PORTRAIT)

    title = document.add_heading(f"{content.target_table} ETL specification", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.style.font.name = "Arial"
    title.style.font.color.rgb = RGBColor.from_string(_INK)
    metadata = document.add_paragraph()
    metadata.add_run(f"OMOP CDM {content.cdm_version}  ·  ").bold = True
    metadata.add_run(f"Sources: {', '.join(content.source_models)}  ·  ")
    metadata.add_run("Deterministic output; no AI request")

    _add_docx_heading(document, "Overview", 1)
    document.add_paragraph(
        content.table_notes
        or content.table_description
        or "No table-level notes recorded."
    )
    _add_docx_heading(document, "Mapping overview", 1)
    figure = _mapping_figure_images(specs)[0]
    with Image.open(BytesIO(figure)) as figure_image:
        image_width, image_height = figure_image.size
    scale = min(7.15 / image_width, 7.25 / image_height)
    picture = document.add_picture(
        BytesIO(figure),
        width=Inches(image_width * scale),
        height=Inches(image_height * scale),
    )
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER

    landscape_section = document.add_section(WD_SECTION.NEW_PAGE)
    _configure_docx_section(landscape_section, WD_ORIENT.LANDSCAPE)
    landscape_section.footer.is_linked_to_previous = True

    if content.relationships:
        _add_docx_heading(document, "Source relationships", 1)
        for relationship in content.relationships:
            document.add_paragraph(relationship, style="List Bullet")

    _add_docx_heading(document, "Field mapping", 1)
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.autofit = False
    widths = (Inches(1.55), Inches(2.15), Inches(3.85), Inches(2.55))
    headers = ("Destination Field", "Source field", "Logic", "Comment field")
    for index, (cell, label) in enumerate(zip(table.rows[0].cells, headers)):
        cell.width = widths[index]
        _shade_docx_cell(cell, _GREEN_SOFT)
        _set_docx_cell_text(cell, label, bold=True)
        _set_docx_cell_margins(cell, top=160, bottom=160)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_properties.append(OxmlElement("w:tblHeader"))
    _keep_docx_row_together(table.rows[0])
    for row in content.rows:
        table_row = table.add_row()
        _keep_docx_row_together(table_row)
        cells = table_row.cells
        for index, value in enumerate(
            (row.destination_field, row.source_field, row.logic, row.comment)
        ):
            cells[index].width = widths[index]
            _set_docx_cell_text(cells[index], value)
            _set_docx_cell_margins(cells[index])

    _add_docx_heading(document, "Change log", 1)
    change_table = document.add_table(rows=1, cols=3)
    change_table.style = "Table Grid"
    for cell, label in zip(
        change_table.rows[0].cells,
        ("Date", "Description", "Author"),
    ):
        _shade_docx_cell(cell, _GREEN_SOFT)
        _set_docx_cell_text(cell, label, bold=True)
        _set_docx_cell_margins(cell, top=160, bottom=160)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    change_header_properties = change_table.rows[0]._tr.get_or_add_trPr()
    change_header_properties.append(OxmlElement("w:tblHeader"))
    _keep_docx_row_together(change_table.rows[0])
    changes = content.changes or (("—", "No changes recorded.", "—"),)
    for change in changes:
        change_row = change_table.add_row()
        _keep_docx_row_together(change_row)
        cells = change_row.cells
        for cell, value in zip(cells, change):
            _set_docx_cell_text(cell, value)
            _set_docx_cell_margins(cell)


def _add_docx_cover(
    document,
    entries: tuple[tuple[EtlSpecificationDocument, ValidatedSpecs], ...],
    *,
    project_name: str,
    project_description: str,
) -> None:
    """Add a concise portrait project cover before table specifications."""
    title = document.add_heading(project_name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.style.font.name = "Arial"
    title.style.font.color.rgb = RGBColor.from_string(_INK)
    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run("OMOP ETL specification")
    subtitle_run.bold = True
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor.from_string(_GREEN)
    document.add_paragraph(
        project_description or "No project description was provided."
    )

    _add_docx_heading(document, "Purpose", 1)
    document.add_paragraph(_PROJECT_PURPOSE)
    _add_docx_heading(document, "Document scope", 1)
    versions = sorted({content.cdm_version for content, _ in entries})
    metadata = document.add_paragraph()
    metadata.add_run("OMOP CDM version: ").bold = True
    metadata.add_run(", ".join(versions))
    metadata = document.add_paragraph()
    metadata.add_run("Included OMOP tables: ").bold = True
    metadata.add_run(str(len(entries)))
    metadata = document.add_paragraph()
    metadata.add_run("Generation method: ").bold = True
    metadata.add_run("Deterministic; no AI request")
    _add_docx_heading(document, "Included tables", 1)
    for content, _ in entries:
        document.add_paragraph(content.target_table, style="List Bullet")
    document.add_page_break()


def _render_docx_collection(
    entries: tuple[tuple[EtlSpecificationDocument, ValidatedSpecs], ...],
    *,
    document_title: str,
    footer_title: str,
    project_name: str,
    project_description: str,
) -> bytes:
    """Render one Word file containing one section pair per OMOP table."""
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(9.5)
    document.core_properties.title = document_title
    document.core_properties.author = "Harmonia OMOP Agent"
    _configure_docx_section(section, WD_ORIENT.PORTRAIT)
    _add_docx_cover(
        document,
        entries,
        project_name=project_name,
        project_description=project_description,
    )
    for index, (content, specs) in enumerate(entries):
        _append_docx_specification(
            document,
            content,
            specs,
            first=index == 0,
        )
    _add_docx_footer(document, footer_title)

    buffer = BytesIO()
    document.save(buffer)
    return _normalize_docx_archive(buffer.getvalue())


def render_docx(
    content: EtlSpecificationDocument,
    specs: ValidatedSpecs,
    *,
    project_name: str = _DEFAULT_PROJECT_NAME,
    project_description: str = "",
) -> bytes:
    """Render a mixed-orientation Microsoft Word ETL specification."""
    return _render_docx_collection(
        ((content, specs),),
        document_title=f"{project_name} — OMOP ETL specification",
        footer_title=content.target_table,
        project_name=project_name,
        project_description=project_description,
    )


def render_docx_collection(
    entries: tuple[tuple[EtlSpecificationDocument, ValidatedSpecs], ...],
    *,
    project_name: str,
    project_description: str,
) -> bytes:
    """Render multiple OMOP table specifications in one Word file."""
    return _render_docx_collection(
        entries,
        document_title=f"{project_name} — OMOP ETL specification",
        footer_title="OMOP",
        project_name=project_name,
        project_description=project_description,
    )


def _pdf_paragraph(value: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(
        html_escape(value).replace("\n", "<br/>") or "—",
        style,
    )


def _pdf_specification_story(
    content: EtlSpecificationDocument,
    specs: ValidatedSpecs,
    *,
    title_style: ParagraphStyle,
    heading_style: ParagraphStyle,
    body_style: ParagraphStyle,
    small_style: ParagraphStyle,
    header_style: ParagraphStyle,
    start_new_portrait_page: bool,
) -> list:
    """Build the portrait overview and landscape tables for one target."""
    story = []
    if start_new_portrait_page:
        story.extend([NextPageTemplate("portrait"), PageBreak()])

    overview = _pdf_paragraph(
        content.table_notes
        or content.table_description
        or "No table-level notes recorded.",
        body_style,
    )
    figure = _mapping_figure_images(specs)[0]
    image = PdfImage(BytesIO(figure))
    scale = min(
        (185 * mm) / image.imageWidth,
        (180 * mm) / image.imageHeight,
        1,
    )
    image.drawWidth = image.imageWidth * scale
    image.drawHeight = image.imageHeight * scale

    story.extend([
        Paragraph(
            f"{html_escape(content.target_table)} ETL specification",
            title_style,
        ),
        _pdf_paragraph(
            f"OMOP CDM {content.cdm_version} · Sources: "
            f"{', '.join(content.source_models)} · Deterministic output; "
            "no AI request",
            body_style,
        ),
        KeepTogether([
            Paragraph("Overview", heading_style),
            overview,
        ]),
        KeepTogether([
            Paragraph("Mapping overview", heading_style),
            image,
        ]),
        NextPageTemplate("landscape"),
        PageBreak(),
    ])
    if content.relationships:
        story.append(Paragraph("Source relationships", heading_style))
        for relationship in content.relationships:
            story.append(_pdf_paragraph(f"• {relationship}", body_style))

    story.extend([
        CondPageBreak(45 * mm),
        Paragraph("Field mapping", heading_style),
    ])
    table_data = [[
        Paragraph("Destination Field", header_style),
        Paragraph("Source field", header_style),
        Paragraph("Logic", header_style),
        Paragraph("Comment field", header_style),
    ]]
    table_data.extend(
        [
            _pdf_paragraph(row.destination_field, small_style),
            _pdf_paragraph(row.source_field, small_style),
            _pdf_paragraph(row.logic, small_style),
            _pdf_paragraph(row.comment, small_style),
        ]
        for row in content.rows
    )
    mapping_table = LongTable(
        table_data,
        colWidths=(38 * mm, 53 * mm, 92 * mm, 72 * mm),
        repeatRows=1,
        hAlign="LEFT",
        splitByRow=1,
    )
    mapping_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{_GREEN_SOFT}")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{_LINE}")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("VALIGN", (0, 0), (-1, 0), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, 0), 8),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
                (
                    "ROWBACKGROUNDS",
                    (0, 1),
                    (-1, -1),
                    [colors.white, colors.HexColor("#F8FAF9")],
                ),
            ]
        )
    )
    story.append(mapping_table)
    story.extend([
        CondPageBreak(35 * mm),
        Paragraph("Change log", heading_style),
    ])
    changes = content.changes or (("—", "No changes recorded.", "—"),)
    change_data = [[
        Paragraph("Date", header_style),
        Paragraph("Description", header_style),
        Paragraph("Author", header_style),
    ]]
    change_data.extend(
        [_pdf_paragraph(value, small_style) for value in change]
        for change in changes
    )
    change_table = LongTable(
        change_data,
        colWidths=(35 * mm, 170 * mm, 50 * mm),
        repeatRows=1,
        hAlign="LEFT",
        splitByRow=1,
    )
    change_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{_GREEN_SOFT}")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(f"#{_LINE}")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("VALIGN", (0, 0), (-1, 0), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, 0), 8),
                ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
            ]
        )
    )
    story.append(change_table)
    return story


def _render_pdf_collection(
    entries: tuple[tuple[EtlSpecificationDocument, ValidatedSpecs], ...],
    *,
    document_title: str,
    footer_title: str,
    project_name: str,
    project_description: str,
) -> bytes:
    """Render one mixed-orientation PDF containing all selected tables."""
    buffer = BytesIO()
    pdf = BaseDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=12 * mm,
        leftMargin=12 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
        invariant=1,
        title=document_title,
        author="Harmonia OMOP Agent",
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "EtlTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        leading=26,
        textColor=colors.HexColor(f"#{_INK}"),
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "EtlHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        leading=16,
        textColor=colors.HexColor(f"#{_GREEN}"),
        spaceBefore=9,
        spaceAfter=5,
    )
    body_style = ParagraphStyle(
        "EtlBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=11,
        textColor=colors.HexColor(f"#{_INK}"),
    )
    small_style = ParagraphStyle(
        "EtlSmall",
        parent=body_style,
        fontSize=7.5,
        leading=9.5,
    )
    header_style = ParagraphStyle(
        "EtlHeader",
        parent=small_style,
        fontName="Helvetica-Bold",
        alignment=TA_CENTER,
        textColor=colors.HexColor(f"#{_GREEN}"),
    )

    def add_page_footer(canvas, document) -> None:
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor(f"#{_MUTED}"))
        page_width, _ = canvas._pagesize
        canvas.drawRightString(
            page_width - 12 * mm,
            5 * mm,
            f"{footer_title} ETL specification  ·  "
            f"Page {document.page}",
        )
        canvas.restoreState()

    portrait_width, portrait_height = A4
    landscape_width, landscape_height = landscape(A4)
    portrait_frame = Frame(
        12 * mm,
        10 * mm,
        portrait_width - 24 * mm,
        portrait_height - 20 * mm,
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
        id="portrait_frame",
    )
    landscape_frame = Frame(
        12 * mm,
        10 * mm,
        landscape_width - 24 * mm,
        landscape_height - 20 * mm,
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
        id="landscape_frame",
    )
    pdf.addPageTemplates([
        PageTemplate(
            id="portrait",
            frames=[portrait_frame],
            pagesize=A4,
            onPage=add_page_footer,
        ),
        PageTemplate(
            id="landscape",
            frames=[landscape_frame],
            pagesize=landscape(A4),
            onPage=add_page_footer,
        ),
    ])

    versions = sorted({content.cdm_version for content, _ in entries})
    story = [
        Paragraph(html_escape(project_name), title_style),
        Paragraph("OMOP ETL specification", heading_style),
        _pdf_paragraph(
            project_description or "No project description was provided.",
            body_style,
        ),
        Paragraph("Purpose", heading_style),
        _pdf_paragraph(_PROJECT_PURPOSE, body_style),
        Paragraph("Document scope", heading_style),
        _pdf_paragraph(
            f"OMOP CDM version: {', '.join(versions)}\n"
            f"Included OMOP tables: {len(entries)}\n"
            "Generation method: Deterministic; no AI request",
            body_style,
        ),
        Paragraph("Included tables", heading_style),
        *(
            _pdf_paragraph(f"• {content.target_table}", body_style)
            for content, _ in entries
        ),
        NextPageTemplate("portrait"),
        PageBreak(),
    ]
    for index, (content, specs) in enumerate(entries):
        story.extend(
            _pdf_specification_story(
                content,
                specs,
                title_style=title_style,
                heading_style=heading_style,
                body_style=body_style,
                small_style=small_style,
                header_style=header_style,
                start_new_portrait_page=index > 0,
            )
        )
    pdf.build(story)
    return buffer.getvalue()


def render_pdf(
    content: EtlSpecificationDocument,
    specs: ValidatedSpecs,
    *,
    project_name: str = _DEFAULT_PROJECT_NAME,
    project_description: str = "",
) -> bytes:
    """Render a mixed-orientation A4 PDF ETL specification."""
    return _render_pdf_collection(
        ((content, specs),),
        document_title=f"{project_name} — OMOP ETL specification",
        footer_title=content.target_table,
        project_name=project_name,
        project_description=project_description,
    )


def render_pdf_collection(
    entries: tuple[tuple[EtlSpecificationDocument, ValidatedSpecs], ...],
    *,
    project_name: str,
    project_description: str,
) -> bytes:
    """Render multiple OMOP table specifications in one PDF file."""
    return _render_pdf_collection(
        entries,
        document_title=f"{project_name} — OMOP ETL specification",
        footer_title="OMOP",
        project_name=project_name,
        project_description=project_description,
    )


def build_etl_specification(
    specs: ValidatedSpecs,
    output_format: EtlSpecificationFormat,
    *,
    project_name: str = _DEFAULT_PROJECT_NAME,
    project_description: str = "",
) -> EtlSpecificationArtifact:
    """Render one validated mapping as Markdown, Word, or PDF."""
    document = build_etl_specification_document(specs)
    entries = ((document, specs),)
    renderers = {
        "md": (
            render_markdown_collection,
            "text/markdown; charset=utf-8",
        ),
        "docx": (
            render_docx_collection,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        "pdf": (render_pdf_collection, "application/pdf"),
    }
    renderer, media_type = renderers[output_format]
    content = renderer(
        entries,
        project_name=project_name.strip() or _DEFAULT_PROJECT_NAME,
        project_description=project_description.strip(),
    )
    if not content:
        raise ValueError("The ETL specification is empty.")
    return EtlSpecificationArtifact(
        content=content,
        file_name=f"{document.target_table}_etl_specification.{output_format}",
        media_type=media_type,
    )


def build_etl_specification_bundle(
    specifications: list[ValidatedSpecs],
    output_format: EtlSpecificationFormat,
    *,
    project_name: str = _DEFAULT_PROJECT_NAME,
    project_description: str = "",
) -> EtlSpecificationArtifact:
    """Render multiple validated mappings in one deterministic document."""
    if len(specifications) < 2:
        raise ValueError("An ETL specification bundle requires two tables.")

    entries = tuple(
        sorted(
            (
                (build_etl_specification_document(specs), specs)
                for specs in specifications
            ),
            key=lambda entry: entry[0].target_table,
        )
    )
    renderers = {
        "md": (
            render_markdown_collection,
            "text/markdown; charset=utf-8",
        ),
        "docx": (
            render_docx_collection,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        "pdf": (render_pdf_collection, "application/pdf"),
    }
    renderer, media_type = renderers[output_format]
    content = renderer(
        entries,
        project_name=project_name.strip() or _DEFAULT_PROJECT_NAME,
        project_description=project_description.strip(),
    )
    if not content:
        raise ValueError("The combined ETL specification is empty.")

    return EtlSpecificationArtifact(
        content=content,
        file_name=f"omop_etl_specification.{output_format}",
        media_type=media_type,
    )
